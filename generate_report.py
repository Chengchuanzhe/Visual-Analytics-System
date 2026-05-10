"""
将实验2实验总结报告.md 转换为格式化的 .docx 文件

import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8")
"""
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import json

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_formatted_paragraph(doc, text, style=None, bold=False, font_size=None, alignment=None, font_name=None, color=None, space_after=None, space_before=None):
    """添加格式化段落"""
    p = doc.add_paragraph(style=style)
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if font_size:
        run.font.size = Pt(font_size)
    if font_name:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_paragraph_with_runs(doc, segments, alignment=None, space_after=6):
    """
    添加混合格式段落
    segments: [(text, bold, font_size, font_name, color), ...]
    """
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    for seg in segments:
        text = seg[0]
        bold = seg[1] if len(seg) > 1 else False
        font_size = seg[2] if len(seg) > 2 else None
        color = seg[3] if len(seg) > 3 else None
        run = p.add_run(text)
        if bold:
            run.bold = True
        if font_size:
            run.font.size = Pt(font_size)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        if color:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_code_block(doc, code_text):
    """添加代码块（灰色背景段落）"""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Cm(1)
        # 灰色背景
        pPr = p._element.get_or_add_pPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="F5F5F5"/>')
        pPr.append(shd)
        run = p.add_run(line if line else ' ')
        run.font.size = Pt(9)
        run.font.name = 'Consolas'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_table(doc, headers, rows, col_widths=None):
    """添加格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, '2F5496')
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if r % 2 == 1:
                set_cell_shading(cell, 'F2F2F2')

    # Set column widths
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()  # spacing after table
    return table


def generate_docx(md_path, output_path):
    doc = Document()

    # ===== 页面设置 =====
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # ===== 样式设置 =====
    style = doc.styles['Normal']
    style.font.size = Pt(12)
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5

    # ===== 封面 / 标题区 =====
    doc.add_paragraph()  # 空行
    add_formatted_paragraph(doc, '中国海洋大学', bold=True, font_size=22, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name='黑体', space_after=8)
    add_formatted_paragraph(doc, '《可视化技术》2026春季实验报告', bold=True, font_size=16, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_name='黑体', space_after=20)

    # 封面信息表
    info_items = [
        ('实验编号：', '实验2'),
        ('实验名称：', '全球海洋-大气耦合时空可视分析系统'),
        ('姓名：', '王哲  许一凡  刘国宁  程传哲'),
        ('专业班级：', '大数据科学与技术  软件工程'),
        ('指导老师：', '解翠'),
        ('实验日期：', '2026年 4月 16日 - 2026年 5月 10日'),
        ('提交时间：', '2026年 5月 10日'),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        run1 = p.add_run(label)
        run1.bold = True
        run1.font.size = Pt(14)
        run1.font.name = '宋体'
        run1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run2 = p.add_run(value)
        run2.font.size = Pt(14)
        run2.font.name = '宋体'
        run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()

    # ==== 读取 Markdown 并逐段处理 ====
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Remove the YAML front matter section (from start to first ---)
    # Actually the md file doesn't have YAML, it starts directly with the report header
    # We already handled the header above, so skip until after the header info line

    # Split by sections
    lines = content.split('\n')

    # We'll parse the markdown programmatically
    i = 0
    while i < len(lines):
        line = lines[i]

        # Skip the header part we already processed
        if line.startswith('**实验编号**') or line.startswith('**实验名称**') or \
           line.startswith('**姓名**') or line.startswith('**专业班级**') or \
           line.startswith('**指导老师**') or line.startswith('**实验日期**') or \
           line.startswith('**提交时间**') or line.startswith('# 全球海洋-大气耦合可视分析系统 · 实验总结报告'):
            i += 1
            continue

        # H1 heading
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            add_formatted_paragraph(doc, text, bold=True, font_size=18, font_name='黑体', space_before=18, space_after=10)
            i += 1
            continue

        # H2 heading
        if line.startswith('## '):
            text = line[3:].strip()
            add_formatted_paragraph(doc, text, bold=True, font_size=15, font_name='黑体', space_before=14, space_after=8)
            i += 1
            continue

        # H3 heading
        if line.startswith('### '):
            text = line[4:].strip()
            add_formatted_paragraph(doc, text, bold=True, font_size=13, font_name='黑体', space_before=10, space_after=6)
            i += 1
            continue

        # H4 heading
        if line.startswith('#### '):
            text = line[5:].strip()
            add_formatted_paragraph(doc, text, bold=True, font_size=12, font_name='黑体', space_before=8, space_after=4)
            i += 1
            continue

        # Code block
        if line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            if code_lines:
                add_code_block(doc, '\n'.join(code_lines))
            continue

        # Table detection: look for the table header pattern "| ... |"
        if line.startswith('|') and '|' in line[1:]:
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1

            # Parse table
            if len(table_lines) >= 2:
                # First line is header, check if second line is separator (|---|---|)
                header_line = table_lines[0]
                sep_line = table_lines[1]
                is_sep = bool(re.match(r'^\|[\s\-:|]+\|$', sep_line))

                if is_sep:
                    headers = [h.strip() for h in header_line.split('|')[1:-1]]
                    rows = []
                    for tl in table_lines[2:]:
                        cells = [c.strip() for c in tl.split('|')[1:-1]]
                        if cells:
                            rows.append(cells)
                    if headers:
                        add_table(doc, headers, rows)
                else:
                    # No separator, treat all as data rows
                    headers = [h.strip() for h in header_line.split('|')[1:-1]]
                    rows = []
                    for tl in table_lines[1:]:
                        cells = [c.strip() for c in tl.split('|')[1:-1]]
                        if cells:
                            rows.append(cells)
                    if headers:
                        add_table(doc, headers, rows)
            continue

        # List items
        if line.startswith('- ') or line.startswith('  - ') or line.startswith('    - '):
            indent_level = (len(line) - len(line.lstrip())) // 2
            text = line.lstrip()[2:]
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(1 + indent_level * 0.8)
            # Bold marker
            bold_pattern = re.match(r'^\*\*(.+?)\*\*(.+)', text)
            if bold_pattern:
                run = p.add_run(bold_pattern.group(1))
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run2 = p.add_run(bold_pattern.group(2))
                run2.font.size = Pt(11)
                run2.font.name = '宋体'
                run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            else:
                run = p.add_run(text)
                run.font.size = Pt(11)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            i += 1
            continue

        # Numbered list
        if re.match(r'^\d+[\.\、]', line):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(line.strip())
            run.font.size = Pt(11)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            i += 1
            continue

        # Horizontal rule / page break
        if line.startswith('---'):
            doc.add_page_break()
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph - handle inline formatting
        text = line.strip()
        if text:
            # Process inline bold, code, etc.
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.first_line_indent = Cm(0.74)

            # Simple inline parsing
            segments = []
            remaining = text
            while remaining:
                # Bold: **text**
                bold_match = re.match(r'^(.*?)\*\*(.+?)\*\*', remaining)
                # Inline code: `text`
                code_match = re.match(r'^(.*?)`(.+?)`', remaining)

                if bold_match and (not code_match or bold_match.start() <= code_match.start()):
                    if bold_match.group(1):
                        segments.append(('normal', bold_match.group(1)))
                    segments.append(('bold', bold_match.group(2)))
                    remaining = remaining[bold_match.end():]
                elif code_match:
                    if code_match.group(1):
                        segments.append(('normal', code_match.group(1)))
                    segments.append(('code', code_match.group(2)))
                    remaining = remaining[code_match.end():]
                else:
                    segments.append(('normal', remaining))
                    remaining = ''

            for seg_type, seg_text in segments:
                run = p.add_run(seg_text)
                if seg_type == 'bold':
                    run.bold = True
                if seg_type == 'code':
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                else:
                    run.font.size = Pt(12)
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        i += 1

    # ===== 保存 =====
    doc.save(output_path)
    print(f'✓ 报告已生成: {output_path}')


if __name__ == '__main__':
    md_path = 'docs/专项文档合集/2-全球海洋大气耦合-实验2-实验总结报告.md'
    output_path = 'docs/专项文档合集/2-全球海洋大气耦合-实验2-实验总结报告.docx'
    generate_docx(md_path, output_path)
